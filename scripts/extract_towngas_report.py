"""Extract the final Towngas screening/pre-FEED report from its working DOCX.

The connected Google Doc contains earlier research notes and drafts before the final
report.  This script keeps only the last report section, beginning at the exact
project title, while preserving the original Word body elements, tables, figures,
styles, and relationships.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_TITLE = (
    "Co-Valorization of Bauxite Residue and Soybean Processing Waste via "
    "Supercritical Water Gasification"
)

TOC_PAGE_NUMBERS = {
    "Abstract": 2,
    "Status, provenance conventions, and design basis": 2,
    "1. Feedstock characterization and blend design": 5,
    "2. Process concept and block architecture": 12,
    "3. Thermodynamic and kinetic basis": 18,
    "4. Detailed process design and equipment definition": 26,
    "5. Closed mass, elemental, and stream balances": 33,
    "6. Energy balance, heat recovery, and Chinese utility basis": 37,
    "7. Heteroatoms, materials, process control, and safety": 39,
    "8. Environmental policy, certification, and product qualification": 41,
    "9. China CAPEX, OPEX, and 20-year economic evaluation in RMB": 44,
    "10. Development programme, risk register, and conclusions": 50,
    "References": 52,
    "Appendices A-D: streams, calculations, economics, and validation register": 53,
}


def _remove_children(parent, tag: str) -> None:
    for child in list(parent):
        if child.tag == qn(tag):
            parent.remove(child)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(95, 101, 98)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _insert_page_break_before(element) -> None:
    """Insert a deterministic hard page break before a body element."""
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    paragraph.append(run)
    element.addprevious(paragraph)


def _format_static_contents(document) -> None:
    """Add deterministic page numbers and dot leaders to the final report TOC."""
    paragraphs = document.paragraphs
    contents_index = next(
        (index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip() == "Contents"),
        None,
    )
    if contents_index is None:
        return

    contents = paragraphs[contents_index]
    contents.paragraph_format.keep_with_next = True
    if contents.runs:
        contents.runs[0].bold = True

    updated = 0
    for paragraph in paragraphs[contents_index + 1 :]:
        if updated == len(TOC_PAGE_NUMBERS):
            break
        label = paragraph.text.strip()
        page_number = TOC_PAGE_NUMBERS.get(label)
        if page_number is None:
            break

        paragraph.text = f"{label}\t{page_number}"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.space_after = Pt(1.5)
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.clear_all()
        tab_stops.add_tab_stop(
            Inches(6.55),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        for run in paragraph.runs:
            run.font.size = Pt(8.5)
        updated += 1


def _move_table_1_5_caption(document) -> None:
    """Place the Table 1.5 caption above its table and keep it with the header."""
    prefix = "Table 1.5 — Co-feed evaluation."
    caption = next(
        (paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith(prefix)),
        None,
    )
    if caption is None:
        return

    table = caption._p.getprevious()
    if table is None or table.tag != qn("w:tbl"):
        return

    table.addprevious(caption._p)
    caption.text = "Table 1.5 — Co-feed evaluation ranked by residue-purity risk."
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.keep_together = True
    caption.paragraph_format.space_after = Pt(2)
    for run in caption.runs:
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(95, 101, 98)


def normalize_layout(document) -> None:
    """Remove inherited working-document layout defects from the report copy."""
    settings = document.settings._element
    _remove_children(settings, "w:mirrorMargins")
    _remove_children(settings, "w:evenAndOddHeaders")

    for section in document.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        section.header_distance = Inches(0.28)
        section.footer_distance = Inches(0.28)
        if section.start_type == WD_SECTION_START.EVEN_PAGE:
            section.start_type = WD_SECTION_START.NEW_PAGE

        footer = section.footer
        footer.is_linked_to_previous = True
        paragraph = footer.paragraphs[0]
        if not paragraph.text:
            _add_page_number(paragraph)

    for paragraph in document.paragraphs:
        if "豆渣" in paragraph.text:
            for run in paragraph.runs:
                if "豆渣" in run.text:
                    run.text = run.text.replace("douzha (豆渣) in Chinese", "douzha in Mandarin")

        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.space_before = Pt(9)
            paragraph.paragraph_format.space_after = Pt(4)

        # These headings inherited hard page starts from the long working
        # document.  In the extracted report they strand most of the preceding
        # page, so let normal pagination place them after the short prior section.
        if paragraph.text in {
            "7. Heteroatoms, materials, process control, and safety",
            "8. Environmental policy, certification, and product qualification",
            "Appendix A - Detailed stream and balance audit",
            "Appendix B - China economic assumptions",
            "Appendix C - Equipment list and calculation basis",
            "Appendix D - Assumption and validation register",
        }:
            paragraph.paragraph_format.page_break_before = False

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True
        properties = table._tbl.tblPr
        _remove_children(properties, "w:tblInd")

        for row_index, row in enumerate(table.rows):
            row_properties = row._tr.get_or_add_trPr()
            _remove_children(row_properties, "w:trHeight")
            if row_properties.find(qn("w:cantSplit")) is None:
                row_properties.append(OxmlElement("w:cantSplit"))
            if row_index == 0 and row_properties.find(qn("w:tblHeader")) is None:
                row_properties.append(OxmlElement("w:tblHeader"))

            # LibreOffice can ignore cantSplit when a long cell paragraph has no
            # paragraph-level pagination controls.  The paired keep_together flag
            # prevents callouts and prose-heavy data rows from breaking mid-line.
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_together = True

        # Table 3.2 is followed by its caption.  Link the final row forward so
        # the caption cannot be orphaned alone at the top of the next page.
        header_text = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
        if header_text.startswith("Co-feed | Case for | Why it is risky"):
            final_cell = table.rows[-1].cells[2]
            final_cell.text = (
                "Public-acceptance risk; compromises ISCC and "
                "fertilizer claims. Rejected."
            )
            for paragraph in final_cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    run.font.size = Pt(8)
        if header_text.startswith("Reaction | Stoichiometry | ΔH°₂₉₈"):
            for cell in table.rows[-1].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True

        # LibreOffice does not reliably honour cantSplit for several prose-heavy
        # rows and one-cell callouts.  Start these compact elements on a fresh
        # page so no heading or first data row is stranded as a fragment.
        break_before_prefixes = (
            "What this decision costs and what it buys",
            "Component | Contributes solids | Contributes carbon",
            "The additive trap:",
            "Calibration warning:",
            "The unifying insight, worth stating explicitly",
            "Property | Screening acceptance | Instrument / response",
        )
        if header_text.startswith(break_before_prefixes):
            target = table._tbl
            if header_text.startswith("Property | Screening acceptance"):
                # Table 4.2 uses a caption above the table; keep it with the table.
                previous = target.getprevious()
                if previous is not None and previous.tag == qn("w:p"):
                    target = previous
            _insert_page_break_before(target)

        # Table 8.2 contains long product-claim text.  Starting the table on a
        # fresh page prevents LibreOffice from splitting a prose-heavy row while
        # using the page that was previously left half empty.
        if header_text.startswith(
            "Stream / claim | Evidence before sale | Base-case economic treatment"
        ):
            target = table._tbl
            previous = target.getprevious()
            if previous is not None and previous.tag == qn("w:p"):
                target = previous
            _insert_page_break_before(target)


def extract_report(source: Path, destination: Path) -> None:
    document = Document(source)
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == REPORT_TITLE]
    if not matches:
        raise RuntimeError(f"Could not find final report title in {source}")

    # The working document contains earlier drafts with the same title.  The final
    # screening/pre-FEED report is the last matching section.
    first_report_paragraph = matches[-1]
    body = document._element.body
    first_element = first_report_paragraph._element

    for element in list(body):
        if element is first_element:
            break
        body.remove(element)

    if "Title" in document.styles:
        first_report_paragraph.style = document.styles["Title"]

    document.core_properties.title = REPORT_TITLE
    document.core_properties.subject = "Final screening/pre-FEED process design report"
    document.core_properties.keywords = (
        "supercritical water gasification, bauxite residue, okara, bi-reforming, OXZEO"
    )

    normalize_layout(document)
    _move_table_1_5_caption(document)
    _format_static_contents(document)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()
    extract_report(args.source, args.destination)


if __name__ == "__main__":
    main()
