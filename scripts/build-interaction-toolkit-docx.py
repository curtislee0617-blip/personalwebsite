from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "curtis-lee-interaction-visualization-toolkit.docx"

NAVY = "203B4A"
GREEN = "466F53"
GREEN_DARK = "31543D"
GREEN_PALE = "E7EEE8"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
BLUE_PALE = "E8EEF5"
CLAY = "B96742"
GOLD = "9A7725"
INK = "26302A"
MUTED = "66726B"
LINE = "CDD6CF"
PAPER = "FBFAF6"
WHITE = "FFFFFF"
CONTENT_DXA = 9360


@dataclass(frozen=True)
class Tool:
    name: str
    role: str
    status: str
    mechanism: str
    ideas: tuple[str, ...]
    pairing: str
    watch: str
    url: str


TOOLS = (
    Tool(
        "Motion",
        "React interface motion",
        "In use",
        "Animates React components, layout changes and gesture-driven values with springs or timelines. Motion values can update every frame without forcing React to re-render.",
        ("Active-stage markers and scientific readouts", "Card/layout continuity on filters", "Subtle entrances that preserve the existing editorial style"),
        "use-gesture, XState, View Transitions",
        "Use it for interface meaning, not for every hover. Respect reduced motion.",
        "https://motion.dev/",
    ),
    Tool(
        "Anime.js",
        "SVG and timeline animation",
        "In use",
        "Targets DOM, CSS, SVG and JavaScript values with compact timelines and stagger utilities. It is especially convenient for path drawing and repeated SVG motion.",
        ("Material-flow dashes in process diagrams", "Draw-on chemistry pathways", "Small diagram sequences with controlled timing"),
        "D3 for geometry; Motion for React state",
        "Assign it one visual property so it does not fight React or another runtime.",
        "https://animejs.com/",
    ),
    Tool(
        "GSAP + ScrollTrigger",
        "Scroll choreography",
        "In use",
        "GSAP supplies a high-control animation engine; ScrollTrigger maps scroll positions to starts, ends, pinning, scrubbing and callbacks.",
        ("Activate B1–B8 as the report column crosses a reading line", "Pinned experimental methods with synchronized figures", "Long-form case-study chapter reveals"),
        "Lenis for a shared frame loop; Motion for UI state",
        "Scope triggers to a component and kill them on unmount. Avoid excessive pinning.",
        "https://gsap.com/docs/v3/Plugins/ScrollTrigger/",
    ),
    Tool(
        "Rive",
        "Interactive vector state machines",
        "Ready",
        "A design-time editor exports compact vector art, animation timelines and state machines. Runtime inputs can change the art in response to application state.",
        ("Purpose-built reactor or separator status artboard", "Coffee grinder cutaway with controllable settings", "Microbe or yeast illustration that responds to selected conditions"),
        "XState for application logic; use-gesture for input",
        "A meaningful .riv asset must be authored first; do not substitute generic stock animation.",
        "https://rive.app/docs/runtimes/react/react/",
    ),
    Tool(
        "dotLottie",
        "Portable authored animation",
        "In use",
        "Packages Lottie vector animation and metadata into a portable file or renders Lottie JSON through a dedicated player. Best for self-contained playback, not complex application logic.",
        ("Material-flow activity indicator", "Compact chapter illustrations", "Loading and empty-state sequences"),
        "After Effects/Lottie authoring; Motion for surrounding UI",
        "Keep essential information in text; animation files still add payload and need reduced-motion behavior.",
        "https://developers.lottiefiles.com/docs/dotlottie-player/dotlottie-react/",
    ),
    Tool(
        "View Transitions API",
        "Native visual continuity",
        "Native + in use",
        "The browser captures old and new visual states and animates between them. Named transition elements can remain visually continuous while the DOM changes.",
        ("Move the SCWG stage marker between B1–B8", "Project-card continuity into a detail page", "Map/list mode changes without a hard visual cut"),
        "Motion for fallback/local animation",
        "Treat it as progressive enhancement and always run the state update when unsupported.",
        "https://developer.mozilla.org/en-US/docs/Web/API/View_Transition_API",
    ),
    Tool(
        "use-gesture",
        "Pointer, drag and pinch input",
        "In use",
        "Normalizes mouse, touch, wheel, drag and pinch into gesture state such as offsets, velocity and intent. It does not render animation by itself.",
        ("Pan and pinch the SCWG flowsheet", "Swipe through image comparisons", "Drag map or flavour objects with clear bounds"),
        "Motion springs or R3F object transforms",
        "Provide keyboard equivalents and set touch-action deliberately.",
        "https://use-gesture.netlify.app/",
    ),
    Tool(
        "XState",
        "State machines and actors",
        "In use",
        "Defines named states, accepted events, legal transitions, context and side effects. The result is an explicit model of what the interface can do, rather than a collection of unrelated booleans.",
        ("Reading vs. inspecting modes in SCWG", "Select / drag equipment / connect stream / validate workflow", "Map loading, filtering, selected-origin and error states"),
        "Motion or Rive can visualize the current state",
        "Use it when behavior is genuinely branching; a single toggle rarely needs a machine.",
        "https://stately.ai/docs",
    ),
    Tool(
        "Lenis",
        "Smooth, synchronized scroll",
        "Ready, opt-in",
        "Keeps native scroll as the source of truth while interpolating an animated position toward the target each frame. It exposes progress and velocity for synchronized DOM or WebGL scenes.",
        ("One immersive project chapter", "Smoother WebGL/DOM synchronization", "Controlled horizontal archive or visual essay"),
        "GSAP ScrollTrigger or WebGL scenes",
        "Do not enable globally first. Test anchors, sticky content, touch, nested scroll and reduced motion.",
        "https://github.com/darkroomengineering/lenis",
    ),
    Tool(
        "D3.js",
        "Custom data and SVG grammar",
        "In use",
        "A modular toolkit for scales, shapes, selections, transitions, geographic projections and data transformation. D3 gives primitives rather than a fixed chart component.",
        ("Existing coffee maps and scientific SVG views", "Custom ternary, phase or process plots", "Geometry for SVG that other engines animate"),
        "Observable Plot for standard charts; React for structure",
        "Keep ownership clear: let React own the DOM unless D3 owns a contained SVG subtree.",
        "https://d3js.org/",
    ),
    Tool(
        "Observable Plot",
        "Exploratory scientific charts",
        "Ready",
        "Builds charts from declarative marks, scales and transforms. It compresses common analytical work into concise code while producing publication-friendly SVG or HTML.",
        ("Roast curves and rate-of-rise", "Extraction yield versus strength", "Temperature, pressure and composition comparisons"),
        "D3 for custom transforms; XState for linked selections",
        "Client plots should be replaced and cleaned up when data changes. Use accessible labels and units.",
        "https://observablehq.com/plot/",
    ),
    Tool(
        "Cytoscape.js",
        "Interactive network visualization",
        "Ready",
        "Represents data as nodes and edges, then applies graph layouts, styles and interaction. Its graph model understands neighborhoods and paths rather than just screen coordinates.",
        ("Coffee variety genealogy", "Flavour-precursor reaction network", "Equipment dependencies and recycle relationships"),
        "XState for guided exploration; Observable Plot for linked metrics",
        "Choose the layout for the question. Large graphs need progressive disclosure and strong labels.",
        "https://js.cytoscape.org/",
    ),
    Tool(
        "MapLibre GL JS",
        "Custom vector maps and globes",
        "Ready",
        "Renders vector tiles with WebGL and a JSON style specification. You control sources, layers, labels, camera, terrain and interaction without tying the site to a proprietary renderer.",
        ("Replace the coffee and wine atlas basemap", "Custom terroir styling and elevation layers", "SCWG site-selection context and infrastructure"),
        "deck.gl for data layers; Turf.js for analysis",
        "The visual style is yours, but tile hosting, attribution and data licenses still matter.",
        "https://maplibre.org/maplibre-gl-js/docs/",
    ),
    Tool(
        "deck.gl",
        "GPU geospatial data layers",
        "Ready",
        "Turns large datasets into WebGL layers such as points, arcs, heatmaps, contours and meshes. A MapboxOverlay can synchronize deck.gl with a MapLibre camera.",
        ("Thousands of coffee farms or tasting pins", "Trade-route arcs and origin-to-roaster journeys", "Animated site constraints, heatmaps and 3D terrain overlays"),
        "MapLibre as basemap; Observable Plot for detail charts",
        "GPU capacity is not a license to show everything. Aggregate and disclose uncertainty.",
        "https://deck.gl/",
    ),
    Tool(
        "Turf.js",
        "Geospatial analysis",
        "Explore next",
        "Applies GIS operations to GeoJSON in the browser or Node: distance, buffer, point-in-polygon, intersections, centroids, grids and more.",
        ("Calculate origin-to-port distance", "Build terroir or facility buffer zones", "Filter farms inside an elevation or administrative polygon"),
        "MapLibre for display; deck.gl for volume",
        "Results inherit the accuracy and projection assumptions of the source data.",
        "https://turfjs.org/",
    ),
    Tool(
        "React Three Fiber",
        "Declarative 3D for React",
        "Ready",
        "A React renderer for Three.js. Cameras, lights, meshes and materials become components while the underlying scene still renders through WebGL.",
        ("Rotate a reactor or roaster cutaway", "3D terrain or equipment explainer", "Molecular geometry where depth improves understanding"),
        "use-gesture, Theatre.js, Lenis",
        "3D must earn its complexity. Load route-locally, cap pixel ratio and provide a 2D fallback.",
        "https://r3f.docs.pmnd.rs/",
    ),
    Tool(
        "PixiJS",
        "High-volume 2D WebGL scenes",
        "Ready",
        "Maintains a GPU-rendered 2D scene graph of sprites, graphics, text and effects. It is a strong bridge between ordinary DOM and full 3D.",
        ("Thousands of aroma or fermentation particles", "A roast-chamber bean field", "Fast illustrated maps or playful recipe transitions"),
        "Matter.js for physics; XState for modes",
        "Canvas content needs an accessible textual counterpart and careful cleanup.",
        "https://pixijs.com/",
    ),
    Tool(
        "Matter.js",
        "Rigid-body physics",
        "In use",
        "Computes gravity, collision, friction, restitution, constraints and pointer interaction for 2D rigid bodies. Rendering can be customized independently of the physics shapes.",
        ("Live falling and draggable coffee beans", "Drag equipment blocks with collisions and snapping", "Ingredient or molecule playgrounds"),
        "PixiJS for dense rendering; XState for tool modes",
        "Physics creates plausible motion, not engineering precision. Use deterministic layout rules for final flowsheets.",
        "https://brm.io/matter-js/",
    ),
    Tool(
        "Theatre.js",
        "Visual animation sequencing",
        "Ready",
        "A timeline editor keyframes JavaScript values and exports project state. Production loads the saved state; the Studio authoring UI should remain development-only.",
        ("Direct a process-chapter hero sequence", "Author a precise 3D camera move", "Coordinate labels, light, camera and material changes"),
        "R3F, PixiJS or DOM/SVG objects",
        "It is a director, not application state. XState should still decide what is legal.",
        "https://www.theatrejs.com/docs/latest",
    ),
    Tool(
        "Tweakpane",
        "Live parameter controls",
        "Explore next",
        "Creates compact control panels for numbers, colors, toggles, folders and monitors. It is ideal while tuning simulations and visual parameters.",
        ("Tune bean gravity, friction and bounce", "Adjust MapLibre/deck layer styling", "Expose educational controls for a reactor or roast model"),
        "Matter.js, R3F, PixiJS, deck.gl",
        "Use it as a lab instrument or development panel; do not let it replace a designed final interface.",
        "https://tweakpane.github.io/docs/",
    ),
    Tool(
        "Tone.js",
        "Interactive web audio",
        "Explore next",
        "Provides a musical timing transport, synthesizers, samplers and effects on top of the Web Audio API. Visual events can be synchronized with sound.",
        ("Optional sonification of a roast curve", "A gentle sound layer for ingredient interactions", "Turn scientific values into pitch or rhythm for accessibility experiments"),
        "XState for playback modes; Motion/Pixi for synchronized visuals",
        "Audio must be opt-in, never autoplay, and include visible controls and a silent equivalent.",
        "https://tonejs.github.io/",
    ),
)


def set_run_font(run, *, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_table_fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(CONTENT_DXA))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[min(index, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))
    set_run_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text, url, *, size=8.5, color=BLUE):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_props.append(run_color)
    run_size = OxmlElement("w:sz")
    run_size.set(qn("w:val"), str(int(size * 2)))
    run_props.append(run_size)
    run_font = OxmlElement("w:rFonts")
    run_font.set(qn("w:ascii"), "Calibri")
    run_font.set(qn("w:hAnsi"), "Calibri")
    run_props.append(run_font)
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_para(
    container,
    text="",
    *,
    size=11,
    color=INK,
    bold=False,
    italic=False,
    before=0,
    after=6,
    align=None,
    keep_next=False,
):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.keep_with_next = keep_next
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_bullet(container, text, *, size=9, color=INK, after=2):
    paragraph = container.add_paragraph()
    paragraph.style = "List Bullet"
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color)
    return paragraph


def add_heading(doc, text, level=1, *, subtitle=None):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=GREEN_DARK, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=GREEN_DARK, bold=True)
    else:
        set_run_font(run, size=12, color=BLUE_DARK, bold=True)
    if subtitle:
        add_para(doc, subtitle, size=9.2, color=MUTED, after=10, keep_next=True)
    return paragraph


def page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    left = hp.add_run("CURTIS LEE  /  INTERACTION & VISUALIZATION TOOLKIT")
    set_run_font(left, size=8, color=MUTED, bold=True)
    right = hp.add_run("                                                        FIELD GUIDE")
    set_run_font(right, size=8, color=GREEN, bold=True)

    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    first_header.paragraphs[0].text = ""

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AUGUST 2026   •   PERSONAL REFERENCE   •   ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(fp)

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = ffp.add_run("CURTIS LEE   •   AUGUST 2026")
    set_run_font(cover_run, size=8.5, color=MUTED, bold=True)

    core_properties = doc.core_properties
    core_properties.title = "Curtis Lee — Interaction & Visualization Toolkit"
    core_properties.subject = "Names, functions, suggested uses, combinations, and implementation notes"
    core_properties.author = "Curtis Lee"
    core_properties.keywords = "animation, visualization, maps, physics, interaction, web"


def add_cover(doc):
    add_para(doc, "CREATIVE TECHNOLOGY FIELD GUIDE", size=9, color=GREEN, bold=True, after=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Interaction &\nVisualization\nToolkit", size=30, color=NAVY, bold=True, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "A practical reference for choosing, combining and imagining animation, mapping, data, physics and state tools across the personal website.",
        size=13,
        color=GREEN_DARK,
        after=38,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    table = doc.add_table(rows=1, cols=3)
    set_table_fixed(table, [2.08, 2.08, 2.08])
    labels = (("21", "tools considered"), ("18", "installed or native"), ("2", "live uses added"))
    for cell, (metric, label) in zip(table.rows[0].cells, labels):
        set_cell_shading(cell, GREEN_PALE)
        set_cell_border(cell, top={"val": "single", "sz": 6, "color": LINE}, bottom={"val": "single", "sz": 6, "color": LINE})
        p1 = add_para(cell, metric, size=22, color=GREEN_DARK, bold=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        p1.paragraph_format.line_spacing = 1
        add_para(cell, label, size=8.5, color=MUTED, bold=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "", size=8, after=22)
    add_para(doc, "CURRENT IMPLEMENTATION", size=8.5, color=CLAY, bold=True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Matter.js coffee-bean physics is live in the coffee guide. XState now governs reading and inspection modes in the gasification process chapter. The remaining large runtimes are route-scoped and opt-in so the earlier website style stays intact.",
        size=10.2,
        color=MUTED,
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def add_choice_page(doc):
    add_heading(doc, "Choose the job before the library", subtitle="The cleanest experiences give each runtime one clear responsibility.")
    rows = (
        ("Animate a React interface", "Motion", "Layout, springs and component state"),
        ("Draw or sequence SVG", "Anime.js", "Paths, timelines and repeated marks"),
        ("Choreograph a scroll chapter", "GSAP + ScrollTrigger", "Triggers, scrub, pin and chapter activation"),
        ("Play authored vector art", "Rive / dotLottie", "Rive for interactive state; dotLottie for playback"),
        ("Show a normal scientific chart", "Observable Plot", "Fast, readable marks from tidy data"),
        ("Invent a bespoke data view", "D3.js", "Scales, shapes, projections and SVG geometry"),
        ("Build a custom geographic atlas", "MapLibre + deck.gl", "Basemap and camera + GPU data layers"),
        ("Explore relationships", "Cytoscape.js", "Nodes, edges, paths and layouts"),
        ("Render many 2D objects", "PixiJS", "GPU sprites, particles and canvas scenes"),
        ("Add collisions or constraints", "Matter.js", "Rigid-body motion and dragging"),
        ("Make legal UI modes explicit", "XState", "States, events, context and actors"),
        ("Direct a precise visual sequence", "Theatre.js", "Keyframes and exported project state"),
    )
    table = doc.add_table(rows=1, cols=4)
    headers = ("If the job is…", "Start with", "If the job is…", "Start with")
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, BLUE_PALE)
        add_para(cell, header, size=9, color=NAVY, bold=True, after=0)
    paired_rows = zip(rows[:6], rows[6:])
    for left, right in paired_rows:
        cells = table.add_row().cells
        values = (left[0], left[1], right[0], right[1])
        for index, (cell, value) in enumerate(zip(cells, values)):
            add_para(cell, value, size=8.3, color=GREEN_DARK if index in (1, 3) else INK, bold=index in (1, 3), after=0)
            set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": LINE})
    set_table_fixed(table, [1.78, 1.2, 1.78, 1.2])
    set_repeat_table_header(table.rows[0])
    add_para(doc, "The later tool cards explain the tradeoffs and the reason behind each default.", size=8.4, color=MUTED, italic=True, after=2)

    add_heading(doc, "Four operating rules", level=2)
    principles = (
        ("One owner per property.", "If Anime.js moves an SVG dash offset, do not let CSS or Motion animate that same offset."),
        ("Progressive enhancement.", "Text, navigation and meaning must remain when motion, WebGL or View Transitions are unavailable."),
        ("Load at the point of use.", "MapLibre, deck.gl, Pixi, R3F and Theatre should stay client-only and route-scoped."),
        ("Scientific restraint.", "Motion should reveal relationships, direction, scale or state—not turn a research page into a product landing page."),
    )
    grid = doc.add_table(rows=2, cols=2)
    set_table_fixed(grid, [3.08, 3.08])
    for cell, (title, body) in zip((cell for row in grid.rows for cell in row.cells), principles):
        set_cell_shading(cell, PAPER)
        set_cell_border(cell, top={"val": "single", "sz": 5, "color": LINE}, bottom={"val": "single", "sz": 5, "color": LINE}, start={"val": "single", "sz": 5, "color": LINE}, end={"val": "single", "sz": 5, "color": LINE})
        add_para(cell, title, size=9, color=GREEN_DARK, bold=True, after=2)
        add_para(cell, body, size=8.4, color=MUTED, after=0)


def add_status_page(doc):
    add_heading(doc, "Toolkit at a glance", subtitle="Status reflects the repository after this implementation pass. “Explore” tools are inspiration additions and are not installed.")
    quick_uses = {
        "Motion": "SCWG UI and site motion",
        "Anime.js": "SCWG material flow",
        "GSAP + ScrollTrigger": "SCWG chapter activation",
        "Rive": "Equipment artboard later",
        "dotLottie": "SCWG activity cue",
        "View Transitions API": "Stage-marker continuity",
        "use-gesture": "SCWG pan and pinch",
        "XState": "Reading / inspection modes",
        "Lenis": "Opt-in wrapper only",
        "D3.js": "Maps and SVG geometry",
        "Observable Plot": "Roast/extraction charts next",
        "Cytoscape.js": "Coffee lineage next",
        "MapLibre GL JS": "Atlas replacement next",
        "deck.gl": "Origin and route layers next",
        "Turf.js": "Map analysis idea",
        "React Three Fiber": "3D scene foundation",
        "PixiJS": "Particle-field foundation",
        "Matter.js": "Live coffee beans",
        "Theatre.js": "Development authoring ready",
        "Tweakpane": "Parameter lab idea",
        "Tone.js": "Optional sonification idea",
    }
    chunks = (TOOLS[:12], TOOLS[12:])
    for chunk_index, chunk in enumerate(chunks):
        if chunk_index:
            heading = add_heading(doc, "Toolkit at a glance — continued", subtitle="Mapping, immersive graphics, physics, sequencing and three extra directions for exploration.")
            heading.paragraph_format.page_break_before = True
        table = doc.add_table(rows=1, cols=4)
        headers = ("Tool", "Primary function", "Status", "Current or first use")
        for cell, header in zip(table.rows[0].cells, headers):
            set_cell_shading(cell, BLUE_PALE)
            add_para(cell, header, size=8.4, color=NAVY, bold=True, after=0)
        for index, tool in enumerate(chunk):
            cells = table.add_row().cells
            values = (tool.name, tool.role, tool.status, quick_uses[tool.name])
            for col, (cell, value) in enumerate(zip(cells, values)):
                if index % 2:
                    set_cell_shading(cell, "F7F8F5")
                color = GREEN_DARK if col == 2 and tool.status != "Explore next" else (CLAY if tool.status == "Explore next" else INK)
                add_para(cell, value, size=7.75, color=color, bold=col in (0, 2), after=0)
                set_cell_border(cell, bottom={"val": "single", "sz": 3, "color": LINE})
        set_table_fixed(table, [1.3, 1.72, 1.0, 2.2])
        set_repeat_table_header(table.rows[0])


def add_tool_card(doc, tool):
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_fixed(table, [3.08, 3.08])

    header = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    set_cell_shading(header, GREEN_PALE if tool.status != "Explore next" else "F7EBDD")
    set_cell_border(header, bottom={"val": "single", "sz": 6, "color": GREEN if tool.status != "Explore next" else CLAY})
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    name = p.add_run(tool.name)
    set_run_font(name, size=12, color=GREEN_DARK if tool.status != "Explore next" else CLAY, bold=True)
    role = p.add_run(f"   {tool.role}")
    set_run_font(role, size=8.5, color=MUTED, bold=True)
    status = p.add_run(f"   •   {tool.status}")
    set_run_font(status, size=8.1, color=CLAY if tool.status == "Explore next" else GREEN, bold=True)

    left, right = table.rows[1].cells
    set_cell_shading(left, WHITE)
    set_cell_shading(right, WHITE)
    add_para(left, "HOW IT WORKS", size=7.8, color=BLUE_DARK, bold=True, after=2)
    add_para(left, tool.mechanism, size=8.55, color=INK, after=0)
    add_para(right, "IDEAS FOR YOUR SITE", size=7.8, color=BLUE_DARK, bold=True, after=2)
    for idea in tool.ideas:
        add_bullet(right, idea, size=8.35, after=1)

    footer = table.rows[2].cells[0].merge(table.rows[2].cells[1])
    set_cell_shading(footer, PAPER)
    p1 = add_para(footer, "", size=8.2, after=2)
    label = p1.add_run("Pairs well: ")
    set_run_font(label, size=8.2, color=GREEN_DARK, bold=True)
    text = p1.add_run(tool.pairing)
    set_run_font(text, size=8.2, color=INK)
    p2 = add_para(footer, "", size=8.2, after=2)
    label = p2.add_run("Watch: ")
    set_run_font(label, size=8.2, color=CLAY, bold=True)
    text = p2.add_run(tool.watch)
    set_run_font(text, size=8.2, color=MUTED)
    source = add_para(footer, "Official documentation  →  ", size=8, color=MUTED, after=0)
    add_hyperlink(source, tool.url.replace("https://", "").rstrip("/"), tool.url, size=8)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.line_spacing = Pt(1)
    spacer.paragraph_format.space_after = Pt(2)
    tiny = spacer.add_run(" ")
    set_run_font(tiny, size=1, color=WHITE)


def add_tool_pages(doc):
    groups = (
        ("Choreography", "Motion, SVG and scroll", ("Motion", "Anime.js", "GSAP + ScrollTrigger")),
        ("Authored assets & continuity", "Interactive vector art, portable playback and native transitions", ("Rive", "dotLottie", "View Transitions API")),
        ("Input & interaction architecture", "Gestures, legal states and scroll synchronization", ("use-gesture", "XState", "Lenis")),
        ("Evidence & relationships", "Custom data grammar, concise plots and graph structure", ("D3.js", "Observable Plot", "Cytoscape.js")),
        ("Maps & spatial reasoning", "Basemap, GPU layers and geospatial analysis", ("MapLibre GL JS", "deck.gl", "Turf.js")),
        ("Scenes, particles & physics", "3D, dense 2D and rigid-body interaction", ("React Three Fiber", "PixiJS", "Matter.js")),
        ("Direction, tuning & sound", "Animation authoring plus three sources of future inspiration", ("Theatre.js", "Tweakpane", "Tone.js")),
    )
    by_name = {tool.name: tool for tool in TOOLS}
    for group_index, (title, subtitle, names) in enumerate(groups):
        if group_index == 0:
            page_break(doc)
        add_heading(doc, title, subtitle=subtitle)
        for name in names:
            add_tool_card(doc, by_name[name])


def add_xstate_page(doc):
    add_heading(doc, "XState: a practical mental model", subtitle="Use a machine when the interface has modes, rules, asynchronous work or states that must never coexist.")

    table = doc.add_table(rows=1, cols=4)
    for cell, label in zip(table.rows[0].cells, ("Concept", "Meaning", "SCWG example", "Equipment editor example")):
        set_cell_shading(cell, BLUE_PALE)
        add_para(cell, label, size=8.5, color=NAVY, bold=True, after=0)
    concepts = (
        ("State", "The current mode", "reading / inspecting", "selecting / dragging / connecting / validating"),
        ("Event", "Something that happened", "INSPECTION.TOGGLE", "EQUIPMENT.DROP or STREAM.START"),
        ("Transition", "A legal response", "reading → inspecting", "connecting → validating"),
        ("Context", "Data carried by the machine", "active B1–B8 ID", "selected unit, stream endpoints"),
        ("Guard", "A condition on a transition", "inspection is available", "target nozzle accepts this stream"),
        ("Actor", "Long-running or async behavior", "future result loader", "validate flowsheet, save layout"),
    )
    for row in concepts:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, row)):
            add_para(cell, text, size=8.2, color=GREEN_DARK if index == 0 else INK, bold=index == 0, after=0)
            set_cell_border(cell, bottom={"val": "single", "sz": 3, "color": LINE})
    set_table_fixed(table, [0.92, 1.38, 1.6, 2.28])
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=35, bottom=35, start=105, end=105)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.05

    add_heading(doc, "Why it is different from several useState calls", level=2)
    add_para(
        doc,
        "Booleans describe facts independently. A state machine describes a valid configuration. With booleans, dragging=true, connecting=true and validating=true can accidentally coexist; with a machine, only declared states and transitions exist. Context stores the selected equipment without pretending every selected item is a new mode.",
        size=9.2,
        color=INK,
        after=5,
    )

    flow = doc.add_table(rows=1, cols=4)
    set_table_fixed(flow, [1.45, 1.45, 1.45, 1.45])
    stages = (("SELECTING", "pointer down"), ("DRAGGING", "drop on canvas"), ("CONNECTING", "choose outlet"), ("VALIDATING", "accept or return"))
    for index, (cell, (state, event)) in enumerate(zip(flow.rows[0].cells, stages)):
        set_cell_shading(cell, GREEN_PALE if index != 3 else BLUE_PALE)
        set_cell_margins(cell, top=40, bottom=40, start=70, end=70)
        add_para(cell, state, size=8.1, color=GREEN_DARK, bold=True, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para(cell, event, size=7.8, color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Good future uses", level=2)
    future = doc.add_table(rows=2, cols=2)
    set_table_fixed(future, [3.08, 3.08])
    future_items = (
        ("Flowsheet editor", "Selection, dragging, stream creation, validation, saving and error recovery."),
        ("Map explorer", "Idle, loading, loaded, filtered, selected and failed states; loaders as actors."),
        ("Recipe laboratory", "Editing, calculating, invalid input, saved and shareable states without contradictory controls."),
        ("Rive artboard", "XState owns application logic; Rive inputs express that state visually."),
    )
    for cell, (title, body) in zip((cell for row in future.rows for cell in row.cells), future_items):
        set_cell_shading(cell, PAPER)
        set_cell_margins(cell, top=45, bottom=45, start=90, end=90)
        set_cell_border(cell, top={"val": "single", "sz": 4, "color": LINE}, bottom={"val": "single", "sz": 4, "color": LINE}, start={"val": "single", "sz": 4, "color": LINE}, end={"val": "single", "sz": 4, "color": LINE})
        add_para(cell, title, size=8.3, color=GREEN_DARK, bold=True, after=1)
        add_para(cell, body, size=8, color=MUTED, after=0)


def add_lenis_page(doc):
    heading = add_heading(doc, "Lenis: what “smooth scroll” actually means", subtitle="The useful part is synchronization and controllable feel—not merely slower scrolling.")
    heading.paragraph_format.page_break_before = True

    model = doc.add_table(rows=1, cols=3)
    set_table_fixed(model, [2.03, 2.03, 2.03])
    steps = (
        ("1  ACTUAL", "The browser's real scroll position remains the foundation."),
        ("2  TARGET", "Wheel, trackpad, touch or scrollTo produces a desired position."),
        ("3  ANIMATED", "Lenis interpolates toward the target on each animation frame and reports progress/velocity."),
    )
    for cell, (title, body) in zip(model.rows[0].cells, steps):
        set_cell_shading(cell, GREEN_PALE)
        set_cell_border(cell, top={"val": "single", "sz": 5, "color": LINE}, bottom={"val": "single", "sz": 5, "color": LINE})
        set_cell_margins(cell, top=45, bottom=45, start=90, end=90)
        add_para(cell, title, size=8.5, color=GREEN_DARK, bold=True, after=3)
        paragraph = add_para(cell, body, size=8.2, color=INK, after=0)
        paragraph.paragraph_format.line_spacing = 1.05

    add_heading(doc, "The controls that matter", level=2)
    table = doc.add_table(rows=1, cols=3)
    for cell, label in zip(table.rows[0].cells, ("Option", "What it changes", "Recommendation here")):
        set_cell_shading(cell, BLUE_PALE)
        add_para(cell, label, size=8.5, color=NAVY, bold=True, after=0)
    options = (
        ("lerp", "How quickly animated scroll catches the target", "Start near 0.1; compare against native"),
        ("autoRaf", "Runs Lenis on requestAnimationFrame", "Enabled inside the local wrapper"),
        ("anchors", "Restores anchor-link handling", "Enabled; verify fixed-header offset"),
        ("allowNestedScroll", "Lets nested scrollers remain native", "Prefer targeted data-lenis-prevent first"),
        ("smoothWheel", "Smooths wheel/trackpad input", "Enabled for the experiment"),
        ("syncTouch", "Adds touch inertia synchronization", "Leave off until iOS testing"),
    )
    for row in options:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, row)):
            paragraph = add_para(cell, text, size=8.05, color=GREEN_DARK if index == 0 else INK, bold=index == 0, after=0)
            paragraph.paragraph_format.line_spacing = 1.05
            set_cell_border(cell, bottom={"val": "single", "sz": 3, "color": LINE})
    set_table_fixed(table, [1.15, 2.28, 2.72])
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=32, bottom=32, start=95, end=95)

    add_heading(doc, "How to use it safely", level=2)
    checklist = (
        "Start on one immersive chapter, not the root layout. Compare the route with Lenis disabled.",
        "Return to native scrolling when prefers-reduced-motion is active.",
        "Test keyboard navigation, browser find, anchor links, sticky figures, nested panels, touch and back/forward restoration.",
        "When pairing with ScrollTrigger, drive Lenis from GSAP's ticker and call ScrollTrigger.update on Lenis scroll events.",
        "Do not use smooth scrolling to conceal an overlong page or an unclear hierarchy.",
    )
    for item in checklist:
        bullet = add_bullet(doc, item, size=8.7, after=1)
        bullet.paragraph_format.line_spacing = 1.05

    note = doc.add_table(rows=1, cols=1)
    set_table_fixed(note, [6.2])
    set_cell_shading(note.cell(0, 0), "F7EBDD")
    set_cell_margins(note.cell(0, 0), top=55, bottom=55, start=100, end=100)
    add_para(note.cell(0, 0), "Recommendation", size=8, color=CLAY, bold=True, after=1)
    add_para(note.cell(0, 0), "Keep Lenis ready but not globally active. Introduce it only if a side-by-side test improves synchronization without harming reading comfort.", size=8.5, color=INK, after=0)


def add_stack_page(doc):
    heading = add_heading(doc, "Combination recipes", subtitle="A stack is strong when each tool solves a different layer of the same experience.")
    heading.paragraph_format.page_break_before = True
    stacks = (
        (
            "1  Scientific SCWG chapter — active now",
            "XState decides reading/inspection state → GSAP detects chapter position → Motion renders progress and controls → D3 frames the SVG → Anime.js moves material streams → use-gesture supplies inspection → dotLottie signals activity → View Transitions preserve deliberate jumps.",
            "Next additions: Observable Plot for result curves; MapLibre + deck.gl for the siting map. Keep 3D optional.",
        ),
        (
            "2  Coffee & wine origin atlas",
            "MapLibre owns the cartographic style and camera → Turf.js calculates distance, buffers and containment → deck.gl draws dense farms, arcs, heatmaps and terrain data → Observable Plot explains the selected place with climate/elevation charts → XState coordinates loading and selection.",
            "This is the highest-value next implementation because the user already wants deeper map customization.",
        ),
        (
            "3  Coffee genealogy explorer",
            "Cytoscape.js positions varieties and relationships → XState holds exploration mode and selected lineage → Observable Plot shows yield, altitude or disease-resistance evidence → Motion transitions the supporting notes.",
            "Keep the graph focused: begin with a lineage, then disclose neighbors instead of showing every cultivar at once.",
        ),
        (
            "4  Draggable equipment laboratory",
            "XState owns select/drag/connect/validate modes → Matter.js handles temporary collision and snapping → PixiJS or SVG renders equipment → Tweakpane exposes engineering/demo parameters → Theatre.js can direct a guided tutorial sequence.",
            "Physics is for interaction feel. Final PFD placement and pipe routing should still follow deterministic engineering rules.",
        ),
        (
            "5  Immersive process or roasting cutaway",
            "React Three Fiber renders the model → Theatre.js authors camera/light/material keyframes → use-gesture allows rotation and inspection → Lenis or ScrollTrigger synchronizes the chapter → XState coordinates labels and modes.",
            "Offer a labelled 2D figure beside the 3D scene so the explanation never depends on WebGL.",
        ),
    )
    for stack_index, (title, recipe, note) in enumerate(stacks):
        if stack_index == 3:
            page_break(doc)
            add_heading(doc, "Combination recipes — continued", subtitle="Two richer experiments for process design and immersive technical explanation.")
        table = doc.add_table(rows=2, cols=1)
        set_table_fixed(table, [6.2])
        set_cell_shading(table.cell(0, 0), GREEN_PALE)
        add_para(table.cell(0, 0), title, size=10, color=GREEN_DARK, bold=True, after=0)
        add_para(table.cell(1, 0), recipe, size=9, color=INK, after=4)
        p = add_para(table.cell(1, 0), "", size=8.6, after=0)
        lead = p.add_run("Design note: ")
        set_run_font(lead, size=8.6, color=CLAY, bold=True)
        rest = p.add_run(note)
        set_run_font(rest, size=8.6, color=MUTED)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_roadmap_page(doc):
    heading = add_heading(doc, "A sensible experimentation roadmap", subtitle="Prioritize ideas that improve the scientific reading experience or an existing map before adding spectacle.")
    heading.paragraph_format.page_break_before = True

    columns = doc.add_table(rows=1, cols=3)
    set_table_fixed(columns, [2.03, 2.03, 2.03])
    roadmap = (
        (
            "NOW  /  already implemented",
            GREEN_PALE,
            ("Matter.js falling and draggable beans", "XState SCWG reading/inspection machine", "Client-only adapters for the remaining installed tools", "Toolkit registry on the Website project card"),
        ),
        (
            "NEXT  /  highest value",
            BLUE_PALE,
            ("MapLibre coffee or wine atlas prototype", "deck.gl origin and journey layers", "Observable Plot roast/extraction figures", "Cytoscape coffee variety family tree"),
        ),
        (
            "LATER  /  authored experiments",
            "F7EBDD",
            ("Rive chemical-equipment artboard", "Theatre-directed 3D cutaway", "Pixi aroma/fermentation particle scene", "Tweakpane lab controls or opt-in Tone.js sonification"),
        ),
    )
    for cell, (title, fill, items) in zip(columns.rows[0].cells, roadmap):
        set_cell_shading(cell, fill)
        add_para(cell, title, size=8.6, color=GREEN_DARK if "LATER" not in title else CLAY, bold=True, after=5)
        for item in items:
            add_bullet(cell, item, size=8.5, after=2)

    add_heading(doc, "Prototype acceptance checklist", level=2)
    checks = (
        ("Meaning", "Can a reader name the relationship, state or data pattern the interaction reveals?"),
        ("Fallback", "Does the content still work without WebGL, pointer input, smooth scroll or animation?"),
        ("Performance", "Is the runtime route-scoped, lazy where practical, cleaned up on unmount and capped for mobile?"),
        ("Control", "Can a user pause, reset, navigate by keyboard and avoid unsolicited audio?"),
        ("Style", "Does it still feel like Curtis's existing site and, on research pages, like a scientific paper?"),
        ("Evidence", "Are units, uncertainty, sources and analytical assumptions visible wherever data is shown?"),
    )
    table = doc.add_table(rows=0, cols=2)
    for label, question in checks:
        cells = table.add_row().cells
        set_cell_shading(cells[0], GREEN_PALE)
        add_para(cells[0], label, size=8.6, color=GREEN_DARK, bold=True, after=0)
        add_para(cells[1], question, size=8.7, color=INK, after=0)
        for cell in cells:
            set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": LINE})
    set_table_fixed(table, [1.05, 5.1])

    add_para(doc, "Prepared as a living reference. Update the status matrix when a Ready tool receives a public-facing use.", size=8.6, color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    page_break(doc)
    add_choice_page(doc)
    page_break(doc)
    add_status_page(doc)
    add_tool_pages(doc)
    add_xstate_page(doc)
    add_lenis_page(doc)
    add_stack_page(doc)
    add_roadmap_page(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
