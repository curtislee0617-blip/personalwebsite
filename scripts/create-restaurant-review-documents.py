#!/usr/bin/env python3

import csv
from collections import Counter
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "docs" / "restaurant-imports"
DATE = "2026-07-27"

ALL_CSV = REPORT_DIR / f"all-restaurants-and-categories-{DATE}.csv"
CLOSED_CSV = REPORT_DIR / f"permanently-closed-restaurants-{DATE}.csv"
ALL_DOCX = REPORT_DIR / f"all-restaurants-and-categories-{DATE}.docx"
CLOSED_DOCX = REPORT_DIR / f"permanently-closed-restaurants-{DATE}.docx"

INK = RGBColor(22, 37, 55)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 102, 112)
LIGHT = RGBColor(226, 233, 241)


def set_run_font(run, size, *, bold=False, italic=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text, url):
    if not url or not urlparse(url).scheme:
        return
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), "17")
    run_properties.extend([color, underline, font_size])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, 8, color=MUTED)


def configure_document(doc, running_label):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = BLUE
    heading_1.paragraph_format.space_before = Pt(14)
    heading_1.paragraph_format.space_after = Pt(6)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = doc.styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_2.font.size = Pt(12.5)
    heading_2.font.bold = True
    heading_2.font.color.rgb = BLUE
    heading_2.paragraph_format.space_before = Pt(10)
    heading_2.paragraph_format.space_after = Pt(4)
    heading_2.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run(running_label), 8, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_masthead(doc, title, subtitle, summary_rows):
    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(8)
    title_paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(title_paragraph.add_run(title), 22, bold=True, color=INK)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.paragraph_format.space_after = Pt(12)
    set_run_font(subtitle_paragraph.add_run(subtitle), 10.5, color=MUTED)

    table = doc.add_table(rows=len(summary_rows), cols=2)
    set_table_geometry(table, [2100, 7260])
    for index, (label, value) in enumerate(summary_rows):
        label_cell, value_cell = table.rows[index].cells
        shade_cell(label_cell, "E8EEF5")
        label_p = label_cell.paragraphs[0]
        value_p = value_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        value_p.paragraph_format.space_after = Pt(0)
        set_run_font(label_p.add_run(label), 9, bold=True, color=INK)
        set_run_font(value_p.add_run(str(value)), 9, color=INK)


def add_restaurant_entry(doc, row, include_status=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    set_run_font(paragraph.add_run(row["name"]), 8.5, bold=True, color=INK)

    details = []
    if include_status:
        details.append(row.get("status", "").replace("_", " ").title())
    if row.get("price_label") and row["price_label"] != "Not available":
        details.append(row["price_label"])
    if row.get("location"):
        details.append(row["location"])
    if row.get("primary_type"):
        details.append(row["primary_type"].replace("_", " "))
    if row.get("source_lists"):
        details.append(row["source_lists"])
    if details:
        set_run_font(paragraph.add_run("  ·  " + "  ·  ".join(details)), 8.5, color=MUTED)
    if row.get("review_reason"):
        set_run_font(paragraph.add_run("  —  " + row["review_reason"]), 8.5, italic=True, color=MUTED)
    if row.get("google_maps_url"):
        paragraph.add_run("  ·  ")
        add_hyperlink(paragraph, "Google Maps", row["google_maps_url"])


def read_csv(path):
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def create_all_restaurants():
    rows = read_csv(ALL_CSV)
    ready = [row for row in rows if row["status"] == "ready"]
    unresolved = [row for row in rows if row["status"] == "needs_review"]
    category_counts = Counter(row["category"] for row in ready)

    doc = Document()
    configure_document(doc, "Restaurant library review · Google Maps Takeout")
    add_masthead(
        doc,
        "Restaurants and categories",
        f"Google Maps Takeout · {DATE} · review copy",
        [
            ("Operational restaurants", f"{len(ready):,}"),
            ("Categories", f"{len(category_counts):,}"),
            ("Manual review", f"{len(unresolved):,} anonymous or unresolved pins"),
            ("Price method", "Google Places where available; category estimate otherwise"),
        ],
    )

    doc.add_heading("Category index", level=1)
    index_table = doc.add_table(rows=0, cols=2)
    set_table_geometry(index_table, [7200, 2160])
    for category, count in sorted(category_counts.items()):
        cells = index_table.add_row().cells
        set_run_font(cells[0].paragraphs[0].add_run(category), 9, bold=True, color=INK)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(cells[1].paragraphs[0].add_run(f"{count:,}"), 9, color=MUTED)
    for row in index_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for category in sorted(category_counts):
        doc.add_heading(f"{category} ({category_counts[category]:,})", level=1)
        for row in (item for item in ready if item["category"] == category):
            add_restaurant_entry(doc, row)

    if unresolved:
        doc.add_heading("Manual review required", level=1)
        note = doc.add_paragraph(
            "These entries could not be identified reliably from their Takeout records, "
            "so they were not guessed or published."
        )
        note.paragraph_format.space_after = Pt(6)
        for row in unresolved:
            add_restaurant_entry(doc, row, include_status=True)

    doc.save(ALL_DOCX)


def create_closed_restaurants():
    rows = read_csv(CLOSED_CSV)
    category_counts = Counter(row["category"] or "Unclassified" for row in rows)

    doc = Document()
    configure_document(doc, "Permanently closed restaurants · Google Maps Takeout")
    add_masthead(
        doc,
        "Permanently closed restaurants",
        f"Google Maps Takeout · {DATE} · verified with Google Places",
        [
            ("Permanently closed", f"{len(rows):,}"),
            ("Temporary closures", "41 excluded from publication but not listed here"),
            ("Website treatment", "Excluded from the published restaurant map"),
            ("Verification", "Official Google Places business status"),
        ],
    )

    doc.add_heading("Category summary", level=1)
    summary_table = doc.add_table(rows=0, cols=2)
    set_table_geometry(summary_table, [7200, 2160])
    for category, count in sorted(category_counts.items()):
        cells = summary_table.add_row().cells
        set_run_font(cells[0].paragraphs[0].add_run(category), 9, bold=True, color=INK)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(cells[1].paragraphs[0].add_run(str(count)), 9, color=MUTED)
    for row in summary_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for category in sorted(category_counts):
        category_rows = [row for row in rows if (row["category"] or "Unclassified") == category]
        doc.add_heading(f"{category} ({len(category_rows)})", level=1)
        for row in category_rows:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.keep_together = True
            set_run_font(paragraph.add_run(row["name"]), 9, bold=True, color=INK)
            details = [row["location"], row["source_lists"]]
            details = [value for value in details if value]
            if details:
                set_run_font(paragraph.add_run("  ·  " + "  ·  ".join(details)), 9, color=MUTED)
            if row["google_maps_url"]:
                paragraph.add_run("  ·  ")
                add_hyperlink(paragraph, "Google Maps", row["google_maps_url"])

    doc.save(CLOSED_DOCX)


if __name__ == "__main__":
    create_all_restaurants()
    create_closed_restaurants()
    print(ALL_DOCX)
    print(CLOSED_DOCX)
